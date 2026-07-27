"""
farmaura-api/app/services/purchase_quote_ai_service.py

Purchase quote AI import service for Farmaura.

Responsibilities:
- extract structured purchase quote data (supplier, payment terms, items,
  freight, delivery time) from a supplier document using AI;
- support PDF/PNG/JPEG via the multimodal document prompt, and XLSX/DOCX by
  parsing them locally first (no multimodal support exists for those formats)
  and feeding the extracted text/table content through a text prompt;
- persist the human-reviewed result as a purchase quote, including the
  original source document, without ever touching sellable inventory;

Observations:
- both Gemini and OpenAI accept PDF and image inline;
- XLSX/DOCX text extraction is capped (LOCAL_PARSE_MAX_CHARS) only as a hard
  safety ceiling against pathological files — the cap must stay high enough
  that real supplier price lists (some run 500-1000+ rows) reach the AI whole,
  because the split-and-retry logic below only ever sees text truncation as an
  *AI output* being cut off (is_response_truncated); an *input* silently cut
  by this local cap before the first call would never trigger a split and
  would just drop the tail of the file with no error;
- layouts vary a lot (flat tables, order forms with a two-row header, sheets
  that repeat the same catalog with different columns, tax-rate-by-state
  reference sheets mixed into the same workbook), so the AI still normalizes
  the extracted table/text into the target schema rather than us trying to
  parse columns ourselves — deterministic code only handles what's safe to
  get unambiguously right (merged cells, obviously-irrelevant sheets);
- a document whose extraction would otherwise be truncated by the provider's
  output token limit gets split (PDF by page range, XLSX/DOCX by text line
  range) and retried per chunk, merging results — see the SPLIT-AND-RETRY
  EXTRACTION section below.
"""

from __future__ import annotations

import asyncio
import base64
import json
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.ai_json import is_response_truncated, parse_ai_json_object
from app.core.config import Settings
from app.core.database import SessionFactory
from app.core.file_storage import write_private_file
from app.core.file_validation import validate_quote_upload
from app.core.tenant_context import apply_tenant_context
from app.repositories.purchase_quote_repository import PurchaseQuoteRepository
from app.schemas.auth import TokenSubject
from app.schemas.purchase_quote import (
    PurchaseQuoteBatchImportConfirmItem,
    PurchaseQuoteBatchImportConfirmResponse,
    PurchaseQuoteBatchImportPreviewItem,
    PurchaseQuoteBatchImportPreviewResponse,
    PurchaseQuoteImportConfirmRequest,
    PurchaseQuoteImportPreviewHeaderResponse,
    PurchaseQuoteImportPreviewLineResponse,
    PurchaseQuoteImportPreviewPaymentTermResponse,
    PurchaseQuoteImportPreviewResponse,
    PurchaseQuoteProductCandidateResponse,
    PurchaseQuoteResponse,
)
from app.services.ai_service import AiDocumentExecutionRequest, AiExecutionRequest, AiService
from app.services.purchase_quote_service import PurchaseQuoteService

DOCUMENT_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg"}
# Bounds a batch import to a reasonable number of sequential AI extractions — files are processed
# one at a time (see BATCH (MULTIPLE FILES) IMPORT below), so this is also a rough cap on how long
# a single batch request can run for.
MAX_BATCH_FILES = 10
# Comfortably above any real supplier price list seen so far (largest sample: ~46k chars across 6
# item-bearing sheets after tax-rate sheets are filtered out). A genuinely pathological workbook —
# hundreds of items *and* multiple very wide (40+ column) internal calculation-mirror sheets that
# aren't filtered by _looks_like_tax_rate_sheet — can still hit this ceiling and lose its tail
# silently; known limitation, not solved here (see Modulo_Orcamentos.md).
LOCAL_PARSE_MAX_CHARS = 200_000

# Header keywords that mark a sheet as an ICMS-ST rate-by-state reference table (e.g. "Estado",
# "Aliquota Interestadual", "MVA", "ST Normal") rather than a list of purchasable items — common
# in supplier workbooks alongside the real price sheet. Skipped entirely from extraction: including
# them wastes LOCAL_PARSE_MAX_CHARS budget on 27-row per-state tax matrices and risks the AI
# mistaking a state row (e.g. "SP, 0.04, 0.18, ...") for a product line.
TAX_RATE_SHEET_MARKERS = ("estado", "aliquota", "alíquota", "mva")

PAYMENT_METHOD_ALIASES: dict[str, str] = {
    "pix": "pix",
    "boleto a vista": "boleto_avista",
    "boleto à vista": "boleto_avista",
    "boleto avista": "boleto_avista",
    "boleto a prazo": "boleto_prazo",
    "boleto parcelado": "boleto_prazo",
    "cartao de credito": "cartao_credito",
    "cartão de crédito": "cartao_credito",
    "cartao credito": "cartao_credito",
    "cartao de debito": "cartao_debito",
    "cartão de débito": "cartao_debito",
    "cartao debito": "cartao_debito",
    "consignado": "consignado",
    "dinheiro": "dinheiro",
    "especie": "dinheiro",
    "espécie": "dinheiro",
    "transferencia": "transferencia",
    "transferência": "transferencia",
    "ted": "transferencia",
    "doc": "transferencia",
}


# ============================================================================
# PURCHASE QUOTE AI SERVICE
# ============================================================================


class PurchaseQuoteAiService:
    """Extract and confirm purchase quote imports for the purchasing module."""

    def __init__(self, session: AsyncSession, subject: TokenSubject, settings: Settings) -> None:
        """Store dependencies required for quote extraction and import."""

        self.session = session
        self.subject = subject
        self.settings = settings
        self.repository = PurchaseQuoteRepository(session)
        self.ai_service = AiService(settings)
        self.quote_service = PurchaseQuoteService(session=session, subject=subject)

    async def preview_quote_import(
        self,
        *,
        file: UploadFile,
        provider: str,
        model: str,
    ) -> PurchaseQuoteImportPreviewResponse:
        """Extract purchase quote data from a supplier document and return a review payload."""

        await validate_quote_upload(file, self.settings)
        content = await file.read(self.settings.max_upload_bytes + 1)
        await file.seek(0)
        if not content:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Quote file is empty."
            )
        normalized_provider = (provider or self.settings.ai_default_provider).strip().lower()
        extension = Path(file.filename or "").suffix.lower()
        file_name = file.filename or "orcamento"

        if extension in DOCUMENT_EXTENSIONS:
            parsed, response_provider, response_model = await self._extract_document_payload(
                content=content,
                mime_type=file.content_type or "application/octet-stream",
                file_name=file_name,
                provider=normalized_provider,
                model=model,
            )
        else:
            extracted_text = (
                self._extract_xlsx_text(content)
                if extension == ".xlsx"
                else self._extract_docx_text(content)
            )
            parsed, response_provider, response_model = await self._extract_text_payload(
                extracted_text=extracted_text,
                file_name=file_name,
                provider=normalized_provider,
                model=model,
            )
        raw_header = parsed.get("quote")
        payment_terms_payload = parsed.get("payment_terms") or []
        items_payload = parsed.get("items") or []
        if not isinstance(payment_terms_payload, list) or not isinstance(items_payload, list):
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail="AI quote extraction returned an invalid payload.",
            )
        header_payload: dict[str, object] = raw_header if isinstance(raw_header, dict) else {}

        supplier_name = self._safe_text(header_payload.get("supplier_name"))
        supplier_document = self._safe_text(header_payload.get("supplier_document"))
        matched_supplier = await self.repository.find_supplier_match(
            tenant_id=str(self.subject.tenant_id),
            cnpj=supplier_document,
            name=supplier_name,
        )

        payment_terms = [
            PurchaseQuoteImportPreviewPaymentTermResponse(
                method=self._safe_payment_method(term.get("method")),
                discount_percent=self._safe_optional_decimal(term.get("discount_percent")),
                surcharge_percent=self._safe_optional_decimal(term.get("surcharge_percent")),
                installment_count=self._safe_optional_int(term.get("installment_count")),
                days_to_pay=self._safe_optional_int(term.get("days_to_pay")),
                notes=self._safe_text(term.get("notes")),
            )
            for term in payment_terms_payload
            if isinstance(term, dict)
        ]

        lines: list[PurchaseQuoteImportPreviewLineResponse] = []
        for index, raw_item in enumerate(items_payload, start=1):
            if not isinstance(raw_item, dict):
                continue
            description = self._safe_text(raw_item.get("description"), fallback=f"Item {index}")
            brand_name = self._safe_text(raw_item.get("brand_name"))
            ean_code = self._safe_text(raw_item.get("ean_code"))
            candidates = await self.repository.search_candidate_products(
                tenant_id=str(self.subject.tenant_id),
                query=(description + " " + brand_name).strip(),
                ean_code=ean_code,
                limit=6,
            )
            units_per_package = self._safe_optional_decimal(raw_item.get("units_per_package"))
            if units_per_package is not None and units_per_package <= Decimal("0.00"):
                units_per_package = None
            lines.append(
                PurchaseQuoteImportPreviewLineResponse(
                    line_id=f"line-{index}",
                    description=description,
                    brand_name=brand_name,
                    sku=self._safe_text(raw_item.get("sku")),
                    ean_code=ean_code,
                    unit=self._safe_text(raw_item.get("unit"), fallback="un"),
                    units_per_package=units_per_package,
                    quantity_reference=self._safe_optional_decimal(
                        raw_item.get("quantity_reference")
                    ),
                    unit_price=self._safe_decimal(raw_item.get("unit_price")),
                    ncm_code=self._safe_text(raw_item.get("ncm_code")),
                    ipi_percentage=self._safe_optional_decimal(raw_item.get("ipi_percentage")),
                    icms_st_value=self._safe_optional_decimal(raw_item.get("icms_st_value")),
                    final_unit_price=self._safe_optional_decimal(
                        raw_item.get("final_unit_price")
                    ),
                    is_comodato=self._safe_bool(raw_item.get("is_comodato")),
                    comodato_notes=self._safe_text(raw_item.get("comodato_notes")),
                    match_candidates=[
                        PurchaseQuoteProductCandidateResponse(
                            id=product.id,
                            name=product.name,
                            brand_name=product.brand_name,
                            sku=product.sku,
                            ean_code=product.ean_code,
                        )
                        for product in candidates
                    ],
                )
            )

        return PurchaseQuoteImportPreviewResponse(
            provider=response_provider,
            model=response_model,
            source_file_name=file.filename or "orcamento",
            header=PurchaseQuoteImportPreviewHeaderResponse(
                supplier_name=supplier_name,
                supplier_document=supplier_document,
                matched_supplier_id=matched_supplier.id if matched_supplier is not None else "",
                quote_date=self._safe_text(header_payload.get("quote_date")),
                valid_until=self._safe_text(header_payload.get("valid_until")),
                freight_type=self._safe_freight_type(header_payload.get("freight_type")),
                freight_cost=self._safe_optional_decimal(header_payload.get("freight_cost")),
                delivery_time_days=self._safe_optional_int(
                    header_payload.get("delivery_time_days")
                ),
                notes=self._safe_text(header_payload.get("notes")),
            ),
            payment_terms=payment_terms,
            items=lines,
        )

    async def confirm_quote_import(
        self,
        *,
        payload: PurchaseQuoteImportConfirmRequest,
        file: UploadFile,
    ) -> PurchaseQuoteResponse:
        """Persist the human-reviewed AI extraction result as a confirmed purchase quote."""

        await validate_quote_upload(file, self.settings)
        content = await file.read(self.settings.max_upload_bytes + 1)
        await file.seek(0)
        if not content:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Quote file is empty."
            )
        extension = Path(file.filename or "").suffix.lower()
        storage_key = f"{self.subject.tenant_id}/purchase-quotes/{uuid4()}{extension}"
        await write_private_file(settings=self.settings, storage_key=storage_key, content=content)

        quote = self.quote_service.build_quote_entity(
            payload,
            status_value="confirmed",
            source_provider=payload.source_provider,
            source_model=payload.source_model,
            file_name=file.filename or "",
            content_type=file.content_type or "",
            size_bytes=len(content),
            storage_key=storage_key,
        )
        return await self.quote_service.persist_quote(quote)

    # ========================================================================
    # BATCH (MULTIPLE FILES) IMPORT
    # ========================================================================
    #
    # A batch is several independent single-file imports run concurrently (bounded by
    # MAX_BATCH_FILES) -- one uploaded file always extracts to exactly one quote/supplier, same as
    # preview_quote_import/confirm_quote_import above; a batch never merges two files' items into
    # one quote. Each file succeeds or fails independently so one malformed document doesn't block
    # the rest.
    #
    # AsyncSession is not safe for concurrent use, so the shared request-scoped self.session can't
    # be used by more than one file at a time. Each concurrent task instead opens its own
    # independent session from the module-level SessionFactory, applies tenant/RLS context to it
    # (apply_tenant_context only needs a session + the subject, not any other request state), and
    # runs the existing single-file method against a throwaway PurchaseQuoteAiService scoped to
    # that session -- see database.py for why the engine's pool was sized up to cover this.

    async def _preview_one(
        self, file: UploadFile, provider: str, model: str
    ) -> PurchaseQuoteBatchImportPreviewItem:
        """Extract one file's preview on its own session, never raising."""

        file_name = file.filename or ""
        async with SessionFactory() as task_session:
            await apply_tenant_context(task_session, self.subject)
            task_service = PurchaseQuoteAiService(
                session=task_session, subject=self.subject, settings=self.settings
            )
            try:
                preview = await task_service.preview_quote_import(
                    file=file, provider=provider, model=model
                )
                return PurchaseQuoteBatchImportPreviewItem(
                    file_name=file_name, success=True, preview=preview, error=""
                )
            except HTTPException as error:
                return PurchaseQuoteBatchImportPreviewItem(
                    file_name=file_name, success=False, preview=None, error=str(error.detail)
                )

    async def preview_quote_import_batch(
        self,
        *,
        files: list[UploadFile],
        provider: str,
        model: str,
    ) -> PurchaseQuoteBatchImportPreviewResponse:
        """Extract purchase quote data from several supplier documents, concurrently."""

        self._ensure_batch_size(files)
        results = await asyncio.gather(
            *(self._preview_one(file, provider, model) for file in files)
        )
        return PurchaseQuoteBatchImportPreviewResponse(results=list(results))

    async def _confirm_one(
        self, file: UploadFile, payload: PurchaseQuoteImportConfirmRequest
    ) -> PurchaseQuoteBatchImportConfirmItem:
        """Persist one file's reviewed result on its own session, never raising."""

        file_name = file.filename or ""
        async with SessionFactory() as task_session:
            await apply_tenant_context(task_session, self.subject)
            task_service = PurchaseQuoteAiService(
                session=task_session, subject=self.subject, settings=self.settings
            )
            try:
                quote = await task_service.confirm_quote_import(payload=payload, file=file)
                return PurchaseQuoteBatchImportConfirmItem(
                    file_name=file_name, success=True, quote=quote, error=""
                )
            except HTTPException as error:
                return PurchaseQuoteBatchImportConfirmItem(
                    file_name=file_name, success=False, quote=None, error=str(error.detail)
                )

    async def confirm_quote_import_batch(
        self,
        *,
        files: list[UploadFile],
        payloads: list[PurchaseQuoteImportConfirmRequest],
    ) -> PurchaseQuoteBatchImportConfirmResponse:
        """Persist several human-reviewed AI extraction results, concurrently."""

        self._ensure_batch_size(files)
        if len(files) != len(payloads):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Each file must have a matching review payload.",
            )
        tasks = [
            self._confirm_one(file, payload)
            for file, payload in zip(files, payloads, strict=True)
        ]
        results = await asyncio.gather(*tasks)
        return PurchaseQuoteBatchImportConfirmResponse(results=list(results))

    def _ensure_batch_size(self, files: list[UploadFile]) -> None:
        """Reject a batch with no files or more than MAX_BATCH_FILES."""

        if not files:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="No files were sent."
            )
        if len(files) > MAX_BATCH_FILES:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Send at most {MAX_BATCH_FILES} files per batch.",
            )

    # ========================================================================
    # SPLIT-AND-RETRY EXTRACTION
    # ========================================================================
    #
    # A single large orçamento can produce more line items than fit in one AI
    # response before the provider's output token limit cuts it off. Instead of
    # just failing, both extraction paths below detect that (is_response_truncated)
    # and split the source in half — by PDF page range, or by extracted-text line
    # range — retrying each half independently (recursively, up to a bounded
    # depth) and merging the results, so a document with many items still gets
    # fully read and shown to the user for review. Images can't be split this
    # way (there's no natural boundary to cut along) and fail with a clear
    # message instead — in practice a single photographed/scanned page rarely
    # has enough items to hit the limit.

    async def _extract_document_payload(
        self,
        *,
        content: bytes,
        mime_type: str,
        file_name: str,
        provider: str,
        model: str,
        max_split_depth: int = 2,
    ) -> tuple[dict[str, object], str, str]:
        """Extract structured data from a document, splitting PDFs by page if truncated."""

        response = await self.ai_service.execute_document_prompt(
            AiDocumentExecutionRequest(
                provider=provider,
                model=model,
                prompt=self._quote_extraction_prompt(file_name),
                system_prompt=self._quote_system_prompt(),
                mime_type=mime_type,
                file_name=file_name,
                file_base64=base64.b64encode(content).decode("ascii"),
                temperature=0.1,
                max_output_tokens=8000,
            )
        )
        if not is_response_truncated(response):
            parsed = parse_ai_json_object(response.content, error_context="AI quote extraction")
            return parsed, response.provider, response.model

        if mime_type != "application/pdf" or max_split_depth <= 0:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=(
                    "A resposta da IA foi cortada por exceder o limite de tamanho da resposta, "
                    "mesmo apos dividir o arquivo. Tente um orcamento com menos itens."
                ),
            )

        try:
            reader = PdfReader(BytesIO(content))
            page_count = len(reader.pages)
        except Exception as error:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not read the PDF file to split it for re-processing.",
            ) from error
        if page_count < 2:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=(
                    "A resposta da IA foi cortada por exceder o limite de tamanho da resposta, "
                    "e o PDF tem uma unica pagina — nao e possivel dividir. Tente um arquivo "
                    "com menos itens."
                ),
            )
        midpoint = page_count // 2
        payloads: list[dict[str, object]] = []
        response_provider = response.provider
        response_model = response.model
        for start, end in ((0, midpoint), (midpoint, page_count)):
            writer = PdfWriter()
            for page_index in range(start, end):
                writer.add_page(reader.pages[page_index])
            buffer = BytesIO()
            writer.write(buffer)
            chunk_payload, response_provider, response_model = await self._extract_document_payload(
                content=buffer.getvalue(),
                mime_type=mime_type,
                file_name=file_name,
                provider=provider,
                model=model,
                max_split_depth=max_split_depth - 1,
            )
            payloads.append(chunk_payload)
        return self._merge_extracted_payloads(payloads), response_provider, response_model

    async def _extract_text_payload(
        self,
        *,
        extracted_text: str,
        file_name: str,
        provider: str,
        model: str,
        max_split_depth: int = 2,
    ) -> tuple[dict[str, object], str, str]:
        """Extract structured data from parsed XLSX/DOCX text, splitting by lines if truncated."""

        response = await self.ai_service.execute_prompt(
            AiExecutionRequest(
                provider=provider,
                model=model,
                prompt=self._quote_extraction_prompt(file_name)
                + "\n\nConteudo extraido do arquivo:\n"
                + extracted_text,
                system_prompt=self._quote_system_prompt(),
                temperature=0.1,
                max_output_tokens=8000,
            )
        )
        if not is_response_truncated(response):
            parsed = parse_ai_json_object(response.content, error_context="AI quote extraction")
            return parsed, response.provider, response.model

        lines = extracted_text.split("\n")
        if len(lines) < 4 or max_split_depth <= 0:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=(
                    "A resposta da IA foi cortada por exceder o limite de tamanho da resposta, "
                    "mesmo apos dividir o arquivo. Tente um orcamento com menos itens."
                ),
            )

        # Repeat the first lines (usually the header/supplier section) in every chunk so
        # header fields survive the split even when they'd otherwise only be in chunk 1.
        header_preamble = "\n".join(lines[:15])
        midpoint = len(lines) // 2
        line_chunks = (lines[:midpoint], lines[midpoint:])
        payloads: list[dict[str, object]] = []
        response_provider = response.provider
        response_model = response.model
        for index, chunk_lines in enumerate(line_chunks):
            joined_lines = "\n".join(chunk_lines)
            chunk_text = joined_lines if index == 0 else header_preamble + "\n" + joined_lines
            chunk_payload, response_provider, response_model = await self._extract_text_payload(
                extracted_text=chunk_text,
                file_name=file_name,
                provider=provider,
                model=model,
                max_split_depth=max_split_depth - 1,
            )
            payloads.append(chunk_payload)
        return self._merge_extracted_payloads(payloads), response_provider, response_model

    def _merge_extracted_payloads(self, payloads: list[dict[str, object]]) -> dict[str, object]:
        """Merge chunked extraction payloads into one, concatenating items.

        Header fields and payment terms come from whichever chunk has non-empty
        data for each — only the chunk containing the document's header section
        will have those populated; the rest contribute items only.
        """

        merged_header: dict[str, object] = {}
        merged_payment_terms: list[object] = []
        merged_items: list[object] = []
        for payload in payloads:
            header = payload.get("quote")
            if isinstance(header, dict):
                for key, value in header.items():
                    if value not in (None, "", 0) and not merged_header.get(key):
                        merged_header[key] = value
            payment_terms = payload.get("payment_terms")
            if isinstance(payment_terms, list) and not merged_payment_terms:
                merged_payment_terms = payment_terms
            items = payload.get("items")
            if isinstance(items, list):
                merged_items.extend(items)
        return {
            "quote": merged_header,
            "payment_terms": merged_payment_terms,
            "items": merged_items,
        }

    # ========================================================================
    # LOCAL DOCUMENT PARSING (XLSX / DOCX)
    # ========================================================================

    def _extract_xlsx_text(self, content: bytes) -> str:
        """Extract a plain-text table representation from an XLSX workbook.

        Resolves merged cells (common in supplier quote headers and in tables
        exported from a pivot table, e.g. a supplier/category name merged down
        several rows) by filling every cell in a merged region with its anchor
        (top-left) value before flattening rows to text. Without this, every
        non-anchor cell in a merge reads as blank, silently dropping the label
        exactly on the rows where it isn't repeated. Uses non-read-only mode so
        `merged_cells.ranges` is reliably populated — quote files are small
        enough that the extra memory cost doesn't matter.
        """

        try:
            workbook = load_workbook(BytesIO(content), data_only=True)
        except Exception as error:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not read the XLSX file.",
            ) from error
        lines: list[str] = []
        for sheet in workbook.worksheets:
            non_anchor_to_anchor: dict[tuple[int, int], tuple[int, int]] = {}
            for merged_range in sheet.merged_cells.ranges:
                anchor = (merged_range.min_row, merged_range.min_col)
                for row_index in range(merged_range.min_row, merged_range.max_row + 1):
                    for col_index in range(merged_range.min_col, merged_range.max_col + 1):
                        if (row_index, col_index) != anchor:
                            non_anchor_to_anchor[(row_index, col_index)] = anchor
            anchor_values: dict[tuple[int, int], object] = {}
            sheet_rows: list[list[str]] = []
            for row_index, row in enumerate(sheet.iter_rows(), start=1):
                cells: list[str] = []
                for col_index, cell in enumerate(row, start=1):
                    value = cell.value
                    if value is not None:
                        anchor_values[(row_index, col_index)] = value
                    else:
                        anchor = non_anchor_to_anchor.get((row_index, col_index))
                        if anchor is not None:
                            value = anchor_values.get(anchor)
                    cells.append("" if value is None else str(value))
                if any(cell.strip() for cell in cells):
                    sheet_rows.append(cells)
            if self._looks_like_tax_rate_sheet(sheet_rows):
                lines.append(
                    f"### Planilha: {sheet.title} "
                    "(tabela de aliquotas de ICMS-ST por estado, ignorada — nao contem itens)"
                )
                continue
            lines.append(f"### Planilha: {sheet.title}")
            lines.extend(" | ".join(cells) for cells in sheet_rows)
        text = "\n".join(lines)
        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="The XLSX file has no readable content.",
            )
        return text[:LOCAL_PARSE_MAX_CHARS]

    def _looks_like_tax_rate_sheet(self, sheet_rows: list[list[str]]) -> bool:
        """Detect a per-state ICMS-ST rate reference table (not a list of items).

        Some supplier workbooks bundle a sheet per product line with columns like
        "Estado", "Aliquota Interestadual/Interna", "MVA" — a tax-calculation lookup table, not
        purchase items. Checks only the first few non-empty rows (the header is always near the
        top) for that specific combination, so an actual item sheet that happens to mention a
        state somewhere is never misclassified.
        """

        for cells in sheet_rows[:3]:
            normalized = [cell.strip().lower() for cell in cells]
            has_estado = "estado" in normalized
            has_rate_column = any(
                marker in cell for cell in normalized for marker in TAX_RATE_SHEET_MARKERS[1:]
            )
            if has_estado and has_rate_column:
                return True
        return False

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract paragraph and table text from a DOCX document."""

        try:
            document = DocxDocument(BytesIO(content))
        except Exception as error:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not read the DOCX file.",
            ) from error
        lines: list[str] = [
            paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        text = "\n".join(lines)
        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="The DOCX file has no readable content.",
            )
        return text[:LOCAL_PARSE_MAX_CHARS]

    # ========================================================================
    # PROMPTS
    # ========================================================================

    def _quote_system_prompt(self) -> str:
        """Return the system prompt for purchase quote extraction."""

        return (
            "Voce extrai dados de orcamentos/cotacoes de fornecedores para uma farmacia, a partir "
            "de documentos, planilhas ou imagens. Responda apenas em JSON valido, sem markdown, "
            "sem comentarios e sem texto adicional. Se um campo nao existir, devolva string "
            "vazia, zero, false ou null conforme o tipo."
        )

    def _quote_extraction_prompt(self, file_name: str) -> str:
        """Build the user prompt used for purchase quote extraction."""

        schema = json.dumps(
            {
                "quote": {
                    "supplier_name": "",
                    "supplier_document": "",
                    "quote_date": "",
                    "valid_until": "",
                    "freight_type": "",
                    "freight_cost": 0,
                    "delivery_time_days": 0,
                    "notes": "",
                },
                "payment_terms": [
                    {
                        "method": "pix",
                        "discount_percent": 0,
                        "surcharge_percent": 0,
                        "installment_count": 0,
                        "days_to_pay": 0,
                        "notes": "",
                    }
                ],
                "items": [
                    {
                        "description": "",
                        "brand_name": "",
                        "sku": "",
                        "ean_code": "",
                        "unit": "un",
                        "units_per_package": 0,
                        "quantity_reference": 0,
                        "unit_price": 0,
                        "ncm_code": "",
                        "ipi_percentage": 0,
                        "icms_st_value": 0,
                        "final_unit_price": 0,
                        "is_comodato": False,
                        "comodato_notes": "",
                    }
                ],
            },
            ensure_ascii=True,
        )
        return (
            "Analise o orcamento/cotacao de compra enviado e extraia um JSON com a estrutura "
            "exata: " + schema + ". Use numeros sem simbolos monetarios. quote_date e "
            "valid_until no formato AAAA-MM-DD quando for possivel identificar, senao string "
            "vazia. freight_type deve ser 'FOB', 'CIF' ou string vazia. method de cada forma de "
            "pagamento deve ser um destes valores: pix, boleto_avista, boleto_prazo, "
            "cartao_credito, cartao_debito, consignado, dinheiro, transferencia, outro. "
            "is_comodato indica que o item e um equipamento cedido pelo fornecedor (ex.: "
            "geladeira/freezer de marca) e nao um produto comprado; comodato_notes descreve as "
            "condicoes do comodato quando existirem. units_per_package e quantas unidades de "
            "venda cabem dentro de uma unidade cotada quando ela for uma embalagem (caixa, "
            "fardo, pacote, display, etc.). Extraia de colunas como 'CX', 'Cx.', 'Multiplo', "
            "'Caixa de Embarque' ou de texto como 'CX C/50', 'caixa com 50 unidades', 'Caixa com "
            "12 Unidades', '12 x 120g' (um numero seguido de 'x' antes do tamanho/peso, comum "
            "numa coluna 'Embalagem' — nesse caso o numero antes do 'x' e o units_per_package, "
            "nao o tamanho depois dele). NAO confunda com uma coluna de 'Multiplo de Venda(s)' "
            "ou 'Quantidade Minima de Pedido' — isso e uma regra de quantas unidades minimas o "
            "cliente deve pedir por vez, nao quantas unidades cabem numa caixa; ignore essas "
            "colunas para units_per_package (deixe 0). Deixe 0 tambem quando a unidade ja for a "
            "unidade de venda (un, ampola, frasco) ou quando a informacao nao aparecer. Quando o "
            "preco listado for o preco de uma caixa/pacote inteiro (ex.: catalogo com 'Caixa com "
            "12 Unidades R$ 31,20'), use essa unidade (ex.: 'cx') em vez de 'un', unit_price "
            "igual ao preco da caixa, e units_per_package igual a quantidade de unidades dentro "
            "dela — nao divida o preco. Ignore linhas que sejam apenas resumo/total (ex.: 'TOTAL "
            "<marca>', 'TOTAL DE PRODUTOS', ou uma lista de categorias no topo de um formulario "
            "de pedido com valor zerado ao lado, antes da tabela real de itens comecar) — essas "
            "nao sao itens cotados. Se o mesmo produto (mesmo codigo ou EAN) aparecer em mais de "
            "uma planilha do arquivo com colunas diferentes (ex.: uma aba resumida e uma aba "
            "'Cadastro' mais completa), gere um unico item combinando os dados das duas em vez "
            "de duplicar; para cada campo, prefira o valor da planilha que o preencher com mais "
            "detalhe (ex.: NCM/EAN geralmente estao só na aba de cadastro). ncm_code e o codigo "
            "de classificacao fiscal do produto (coluna 'NCM', "
            "normalmente 8 digitos); deixe vazio se nao aparecer. ipi_percentage e a aliquota "
            "percentual do IPI (coluna '% IPI', 'IPI', 'IPI%'), como numero (ex.: coluna com "
            "'6.5' vira 6.5, nao 0.065); deixe 0 se nao aparecer. icms_st_value e o valor "
            "monetario de substituicao tributaria por unidade (coluna 'ST', 'ICMS-ST', 'Subst. "
            "Tributaria'); deixe 0 se nao aparecer ou se a coluna so tiver zeros. "
            "final_unit_price e o preco unitario final ja com os impostos aplicados, quando o "
            "documento tiver uma coluna propria para isso (ex.: 'Preco Final', 'Preco com "
            "Imposto', 'Valor Final'); nesse caso unit_price continua sendo o preco base (ex.: "
            "coluna 'Preco Unit.') e final_unit_price o valor com imposto — nao calcule "
            "final_unit_price voce mesmo, so preencha quando o documento ja trouxer essa coluna "
            "pronta; deixe 0 quando nao existir. Quando o documento tiver colunas separadas de "
            "codigo do produto e descricao (ex.: 'Codigo' e 'Detalhe', ou 'Codigo' e 'Produto'), "
            "use a coluna de codigo em sku e a coluna de texto/descricao completa em "
            "description. Se o conteudo vier de "
            "uma planilha e parecer "
            "com uma tabela dinamica exportada do Excel — linhas onde a celula de agrupamento "
            "(ex.: nome do fornecedor ou categoria) fica em branco nas linhas seguintes ao "
            "primeiro item do grupo — considere que o valor em branco repete o da linha anterior "
            "daquela mesma coluna. Considere o arquivo " + file_name + "."
        )

    # ========================================================================
    # NORMALIZATION HELPERS
    # ========================================================================

    def _safe_text(self, value: object, *, fallback: str = "") -> str:
        """Normalize optional text values."""

        return str(value or fallback).strip()

    def _safe_bool(self, value: object) -> bool:
        """Normalize optional truthy values to booleans."""

        if isinstance(value, bool):
            return value
        normalized = str(value or "").strip().lower()
        return normalized in {"1", "true", "yes", "sim"}

    def _safe_decimal(self, value: object) -> Decimal:
        """Normalize optional numeric values to non-negative decimals."""

        try:
            normalized = Decimal(str(value or "0")).quantize(Decimal("0.01"))
        except (InvalidOperation, ValueError, TypeError):
            normalized = Decimal("0.00")
        return normalized if normalized >= Decimal("0.00") else Decimal("0.00")

    def _safe_optional_decimal(self, value: object) -> Decimal | None:
        """Normalize an optional numeric value, keeping None when the AI didn't extract it."""

        if value is None or str(value).strip() == "":
            return None
        try:
            normalized = Decimal(str(value)).quantize(Decimal("0.01"))
        except (InvalidOperation, ValueError, TypeError):
            return None
        return normalized if normalized >= Decimal("0.00") else None

    def _safe_optional_int(self, value: object) -> int | None:
        """Normalize an optional integer value, keeping None when the AI didn't extract it."""

        if value is None or str(value).strip() == "":
            return None
        try:
            normalized = int(Decimal(str(value)))
        except (InvalidOperation, ValueError, TypeError):
            return None
        return normalized if normalized >= 0 else None

    def _safe_payment_method(self, value: object) -> str:
        """Normalize a free-text payment method into the supported enum, defaulting to 'outro'."""

        normalized = str(value or "").strip().lower()
        if normalized in PAYMENT_METHOD_ALIASES.values():
            return normalized
        return PAYMENT_METHOD_ALIASES.get(normalized, "outro")

    def _safe_freight_type(self, value: object) -> str:
        """Normalize a free-text freight type into 'FOB', 'CIF', or empty."""

        normalized = str(value or "").strip().upper()
        return normalized if normalized in {"FOB", "CIF"} else ""
